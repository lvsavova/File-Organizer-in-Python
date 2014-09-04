from docx import *


class DocumentHandler:
    """
    A class responsible for searching in the docx files contents
    and organization report file creation. Uses methods from the
    python docx module
    """
    def appears_in_paragraph(self, paragraph, match_string):
        """
        Searches a string in a text paragraph of a docx file.
        Returns True if the string appears in the paragraph.
        It's a helper method for search_content method
        """
        return match_string in paragraph.text

    def search_content(self, document_path, match_string):
        """
        For each paragraph of the given docx document checks
        whether the given string appears in a paragraph.
        Returns True if the string appears at least once.
        Breaks if the string is found
        """
        document = Document(document_path)
        contains_word = False
        for paragraph in document.paragraphs:
            if self.appears_in_paragraph(paragraph, match_string):
                contains_word = True
                break
        return contains_word

    def create_report(self, report_folder, log):
        """
        Creates a docx file in the given folder. The log argument
        is a dictionary with name of file as a key and its destination.
        Writes the log to a word table. Returns a document object
        """
        document = Document()
        document.add_heading('File Organizer Report', 0)
        table = document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'File'
        hdr_cells[1].text = 'Copy Destination'
        for key in log:
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = log[key]
        document.save(report_folder + "\\" + "report.docx")
        return document
